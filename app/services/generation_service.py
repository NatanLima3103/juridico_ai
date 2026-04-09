from datetime import date, datetime, time, timedelta
from io import BytesIO

try:
    from zoneinfo import ZoneInfo, ZoneInfoNotFoundError
except ImportError:
    ZoneInfo = None
    ZoneInfoNotFoundError = Exception

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy import or_
from sqlalchemy.orm import Session, joinedload

from app.models.document import Document
from app.models.generation import Generation
from app.models.writing_profile import WritingProfile
from app.services.ai_generation_service import AIGenerationResult, gerar_resultado_juridico_com_fallback
from app.services.audit_service import registrar_evento_auditoria, serializar_entidade_para_auditoria
from app.services.prompt_service import build_advanced_prompt, build_smart_context


TIPOS_DE_DOCUMENTO = [
    "Petição inicial",
    "Contestação",
    "Réplica",
    "Manifestação",
    "Parecer jurídico",
    "Contrato",
    "Notificação extrajudicial",
    "Recurso",
    "Outro",
]

ORDENACOES_GERACOES = {
    "updated_desc": "AtualizaÃ§Ã£o (mais recente primeiro)",
    "updated_asc": "AtualizaÃ§Ã£o (mais antiga primeiro)",
    "created_desc": "CriaÃ§Ã£o (mais recente primeiro)",
    "created_asc": "CriaÃ§Ã£o (mais antiga primeiro)",
    "client_asc": "Cliente (A-Z)",
    "client_desc": "Cliente (Z-A)",
}


TEMPLATES_JURIDICOS_PRONTOS = {'Petição inicial': {'titulo': 'Modelo base para ajuizamento da ação',
                     'descricao': 'Estrutura pronta para iniciar uma ação com fatos, pedidos e fundamentos essenciais.',
                     'case_subject': 'Ação de obrigação de fazer c/c indenização por danos morais',
                     'facts': 'A parte autora relata que manteve relação contratual com a parte ré e, apesar das '
                              'tentativas administrativas de solução, o problema persistiu. Houve falha na prestação '
                              'do serviço, prejuízos concretos ao cliente e ausência de resposta eficaz no prazo '
                              'razoável. Os documentos anexos demonstram a linha do tempo dos fatos, os protocolos '
                              'realizados e os impactos suportados.',
                     'requests': 'Concessão da tutela provisória, se cabível.\n'
                                 'Citação da parte ré para apresentar resposta.\n'
                                 'Procedência dos pedidos com a condenação da parte ré à obrigação de fazer ou não '
                                 'fazer pertinente ao caso.\n'
                                 'Condenação ao pagamento de indenização pelos prejuízos materiais e/ou morais, quando '
                                 'aplicável.\n'
                                 'Condenação ao pagamento das custas processuais e honorários advocatícios.',
                     'legal_basis': 'Aplicação do Código Civil e/ou do Código de Defesa do Consumidor, conforme a '
                                    'natureza da relação jurídica. Observância dos princípios da boa-fé objetiva, '
                                    'reparação integral do dano, responsabilidade civil e efetividade da tutela '
                                    'jurisdicional. Considerar, ainda, a jurisprudência dominante sobre falha na '
                                    'prestação do serviço e dever de indenizar.'},
 'Contestação': {'titulo': 'Modelo base de defesa',
                 'descricao': 'Estrutura pronta para impugnar fatos, fundamentos e pedidos formulados pela parte '
                              'autora.',
                 'case_subject': 'Ação indenizatória por suposta falha contratual',
                 'facts': 'A parte ré sustenta que os fatos narrados na petição inicial não ocorreram da forma '
                          'descrita pela parte autora. Houve cumprimento das obrigações assumidas, inexistindo conduta '
                          'ilícita, defeito no serviço ou nexo causal apto a justificar a responsabilização '
                          'pretendida. Os documentos juntados demonstram a regularidade da atuação da parte ré e a '
                          'improcedência da narrativa inicial.',
                 'requests': 'Recebimento da presente contestação.\n'
                             'Rejeição do pedido de tutela provisória, se houver.\n'
                             'Total improcedência dos pedidos formulados na inicial.\n'
                             'Produção de todas as provas admitidas em direito.\n'
                             'Condenação da parte autora ao pagamento de custas e honorários advocatícios.',
                 'legal_basis': 'Aplicação do contraditório e da ampla defesa. Ausência dos pressupostos da '
                                'responsabilidade civil, sobretudo ato ilícito, dano comprovado e nexo causal. '
                                'Possibilidade de impugnação específica dos fatos e documentos, bem como incidência '
                                'das regras processuais atinentes ao ônus da prova e à improcedência dos pedidos.'},
 'Réplica': {'titulo': 'Modelo base para impugnação da contestação',
             'descricao': 'Estrutura pronta para rebater preliminares, argumentos defensivos e documentos da parte '
                          'contrária.',
             'case_subject': 'Impugnação aos argumentos apresentados em contestação',
             'facts': 'A contestação apresentada pela parte ré não afasta os fatos narrados na petição inicial. Os '
                      'argumentos defensivos são genéricos, não impugnam de forma específica os elementos centrais da '
                      'demanda e não descaracterizam os prejuízos suportados pela parte autora. Os documentos já '
                      'acostados aos autos permanecem suficientes para demonstrar a plausibilidade da pretensão '
                      'deduzida.',
             'requests': 'Rejeição das preliminares suscitadas na contestação, se existentes.\n'
                         'Acolhimento da impugnação aos argumentos defensivos.\n'
                         'Prosseguimento regular do feito.\n'
                         'Ao final, total procedência dos pedidos iniciais.',
             'legal_basis': 'Aplicação do princípio da dialeticidade processual, da impugnação específica e das regras '
                            'do ônus da prova. Reforço dos fundamentos já apresentados na inicial, com destaque para a '
                            'insuficiência da defesa, manutenção da coerência narrativa e prevalência do conjunto '
                            'probatório já produzido.'},
 'Manifestação': {'titulo': 'Modelo base de manifestação processual',
                  'descricao': 'Estrutura pronta para petições intermediárias, esclarecimentos e pedidos incidentais.',
                  'case_subject': 'Manifestação sobre documento, decisão ou andamento processual',
                  'facts': 'No curso do processo, sobreveio fato processual relevante que exige pronunciamento da '
                           'parte. Trata-se de questão que demanda esclarecimento objetivo, com indicação do contexto, '
                           'do ocorrido e da providência que se busca perante o juízo. A presente manifestação '
                           'delimita o ponto controvertido e apresenta a posição processual da parte interessada.',
                  'requests': 'Recebimento da presente manifestação.\n'
                              'Análise do ponto específico submetido ao juízo.\n'
                              'Adoção da providência processual cabível conforme o caso concreto.\n'
                              'Regular prosseguimento do feito.',
                  'legal_basis': 'Aplicação das normas processuais pertinentes ao ato praticado, observando-se '
                                 'contraditório, cooperação processual, boa-fé e efetividade. Indicar artigos do CPC '
                                 'compatíveis com a providência pretendida, bem como eventuais precedentes ou '
                                 'entendimentos aplicáveis ao tema.'},
 'Parecer jurídico': {'titulo': 'Modelo base de parecer',
                      'descricao': 'Estrutura pronta para análise técnica consultiva de questão jurídica.',
                      'case_subject': 'Análise de viabilidade jurídica da medida pretendida',
                      'facts': 'Foi apresentada consulta jurídica para exame da viabilidade de determinada medida, à '
                               'luz dos fatos narrados e da documentação disponibilizada. O objetivo é avaliar riscos, '
                               'fundamentos jurídicos, alternativas possíveis e providências recomendáveis, '
                               'considerando o cenário fático e normativo aplicável.',
                      'requests': 'Esclarecimento da controvérsia jurídica apresentada.\n'
                                  'Indicação dos fundamentos normativos aplicáveis.\n'
                                  'Avaliação dos riscos e das possíveis estratégias.\n'
                                  'Conclusão objetiva sobre a viabilidade da medida analisada.',
                      'legal_basis': 'Análise sistemática da legislação incidente, da jurisprudência pertinente e dos '
                                     'princípios aplicáveis ao caso. Considerar interpretação literal, teleológica e '
                                     'prática das normas, bem como eventuais entendimentos administrativos ou '
                                     'judiciais relevantes para a conclusão do parecer.'},
 'Contrato': {'titulo': 'Modelo base de minuta contratual',
              'descricao': 'Estrutura pronta para iniciar uma minuta com objeto, obrigações e cláusulas essenciais.',
              'case_subject': 'Prestação de serviços profissionais',
              'facts': 'As partes pretendem formalizar relação contratual para disciplinar prestação de serviços, '
                       'obrigações recíprocas, forma de pagamento, prazo de vigência, hipóteses de rescisão e '
                       'responsabilidades. A minuta deverá refletir com clareza o objeto do ajuste, os limites de '
                       'atuação de cada parte e mecanismos de segurança jurídica para a execução contratual.',
              'requests': 'Definição clara do objeto contratual.\n'
                          'Previsão das obrigações e responsabilidades de cada parte.\n'
                          'Estabelecimento de prazo, remuneração e forma de pagamento.\n'
                          'Definição de cláusulas de rescisão, penalidades e foro.\n'
                          'Adequação da minuta à legislação aplicável.',
              'legal_basis': 'Aplicação dos princípios contratuais, autonomia privada, boa-fé objetiva, função social '
                             'do contrato e força obrigatória das convenções. Observar as regras do Código Civil '
                             'pertinentes ao tipo contratual, além de cláusulas essenciais para segurança, equilíbrio '
                             'e executabilidade do ajuste.'},
 'Notificação extrajudicial': {'titulo': 'Modelo base de notificação',
                               'descricao': 'Estrutura pronta para comunicação formal, constituição em mora e '
                                            'exigência de providência.',
                               'case_subject': 'Constituição em mora por descumprimento contratual',
                               'facts': 'A parte notificante relata inadimplemento ou descumprimento de obrigação '
                                        'assumida pela parte notificada, apesar das tentativas prévias de solução '
                                        'amigável. Diante da persistência do problema, faz-se necessária comunicação '
                                        'formal para ciência inequívoca da irregularidade e fixação de prazo para '
                                        'regularização da conduta.',
                               'requests': 'Ciência formal da irregularidade apontada.\n'
                                           'Regularização da obrigação no prazo assinalado.\n'
                                           'Abstenção de nova conduta lesiva, se aplicável.\n'
                                           'Advertência acerca das medidas judiciais e extrajudiciais cabíveis em caso '
                                           'de inércia.',
                               'legal_basis': 'Aplicação das disposições contratuais pertinentes, dos princípios da '
                                              'boa-fé objetiva e das regras legais sobre mora, inadimplemento e '
                                              'responsabilidade civil. A notificação serve como instrumento formal de '
                                              'constituição em mora e preservação de prova da comunicação realizada.'},
 'Recurso': {'titulo': 'Modelo base recursal',
             'descricao': 'Estrutura pronta para atacar decisão judicial com síntese dos fatos e fundamentos '
                          'recursais.',
             'case_subject': 'Insurgência contra decisão interlocutória ou sentença',
             'facts': 'Foi proferida decisão judicial que causou prejuízo à parte recorrente, exigindo reavaliação '
                      'pelo órgão competente. A decisão impugnada contém equívoco na apreciação dos fatos, na '
                      'interpretação do direito aplicável ou na valoração do conjunto probatório, circunstâncias que '
                      'justificam a interposição do recurso cabível.',
             'requests': 'Recebimento e processamento do recurso.\n'
                         'Reforma ou anulação da decisão recorrida.\n'
                         'Atribuição de efeito suspensivo, se cabível.\n'
                         'Intimação da parte contrária para apresentação de contrarrazões.\n'
                         'Ao final, provimento do recurso.',
             'legal_basis': 'Indicar os pressupostos de admissibilidade recursal, a tempestividade, o cabimento e os '
                            'fundamentos jurídicos para reforma da decisão. Explorar violação legal, error in '
                            'procedendo, error in judicando, descompasso probatório e jurisprudência aplicável ao '
                            'tema.'},
 'Outro': {'titulo': 'Modelo genérico de minuta jurídica',
           'descricao': 'Estrutura flexível para peças não contempladas nas categorias principais.',
           'case_subject': 'Providência jurídica específica do caso concreto',
           'facts': 'Descreva, neste campo, o contexto fático principal, indicando quem são as partes envolvidas, qual '
                    'a controvérsia existente, quais eventos relevantes ocorreram e quais documentos dão suporte à '
                    'narrativa. Organize os fatos em ordem lógica e objetiva para facilitar a redação da minuta.',
           'requests': 'Recebimento da peça apresentada.\n'
                       'Adoção da providência jurídica pretendida.\n'
                       'Produção das provas cabíveis.\n'
                       'Demais requerimentos compatíveis com a natureza da medida.',
           'legal_basis': 'Indique os dispositivos legais, princípios, entendimentos jurisprudenciais e teses '
                          'jurídicas mais relevantes para sustentar a peça. Ajuste o conteúdo conforme o procedimento '
                          'aplicável, a estratégia processual e a natureza do direito discutido.'}}


def listar_templates_juridicos_prontos() -> dict:
    return TEMPLATES_JURIDICOS_PRONTOS


def resumir_texto(texto: str | None, limite: int = 80) -> str:
    texto_limpo = " ".join((texto or "").split())

    if not texto_limpo:
        return "—"

    if len(texto_limpo) <= limite:
        return texto_limpo

    return texto_limpo[:limite].rstrip(" .,;:") + "..."


def coletar_ids_inteiros_unicos(values: list[str] | None) -> list[int]:
    ids: list[int] = []
    for value in values or []:
        try:
            numero = int(value)
        except (TypeError, ValueError):
            continue
        if numero not in ids:
            ids.append(numero)
    return ids


def agora_brasil():
    if ZoneInfo is not None:
        try:
            return datetime.now(ZoneInfo("America/Sao_Paulo")).replace(tzinfo=None)
        except ZoneInfoNotFoundError:
            pass
        except Exception:
            pass

    return datetime.now()


def _parse_date_input(raw_value: str | None) -> date | None:
    valor = (raw_value or "").strip()

    if not valor:
        return None

    try:
        return datetime.strptime(valor, "%Y-%m-%d").date()
    except ValueError:
        return None


def _normalizar_texto(valor: str | None) -> str:
    return (valor or "").strip()


def normalizar_filtros_listagem(
    search_term: str = "",
    document_type: str = "",
    writing_profile_id: int | None = None,
    sem_perfil: bool = False,
    client_name: str = "",
    case_subject: str = "",
    created_from: str = "",
    created_to: str = "",
    sort_by: str = "updated_desc",
) -> dict:
    opcoes_ordenacao = {
        "updated_desc",
        "updated_asc",
        "created_desc",
        "created_asc",
        "client_asc",
        "client_desc",
    }

    return {
        "search_term": _normalizar_texto(search_term),
        "document_type": _normalizar_texto(document_type),
        "writing_profile_id": writing_profile_id if isinstance(writing_profile_id, int) and writing_profile_id > 0 else None,
        "sem_perfil": bool(sem_perfil),
        "client_name": _normalizar_texto(client_name),
        "case_subject": _normalizar_texto(case_subject),
        "created_from": _parse_date_input(created_from),
        "created_to": _parse_date_input(created_to),
        "sort_by": sort_by if sort_by in opcoes_ordenacao else "updated_desc",
    }


def contar_filtros_ativos_geracoes(filtros: dict | None) -> int:
    if not filtros:
        return 0

    total = 0

    if filtros.get("search_term"):
        total += 1
    if filtros.get("document_type"):
        total += 1
    if filtros.get("writing_profile_id"):
        total += 1
    if filtros.get("sem_perfil"):
        total += 1
    if filtros.get("client_name"):
        total += 1
    if filtros.get("case_subject"):
        total += 1
    if filtros.get("created_from"):
        total += 1
    if filtros.get("created_to"):
        total += 1
    if filtros.get("sort_by", "updated_desc") != "updated_desc":
        total += 1

    return total


def serializar_filtros_geracao_para_template(filtros: dict | None) -> dict:
    filtros = filtros or {}

    if filtros.get("sem_perfil"):
        writing_profile_id = "none"
    elif filtros.get("writing_profile_id"):
        writing_profile_id = str(filtros.get("writing_profile_id"))
    else:
        writing_profile_id = ""

    return {
        "search": filtros.get("search_term", ""),
        "search_term": filtros.get("search_term", ""),
        "document_type": filtros.get("document_type", ""),
        "writing_profile_id": writing_profile_id,
        "sem_perfil": bool(filtros.get("sem_perfil")),
        "client_name": filtros.get("client_name", ""),
        "case_subject": filtros.get("case_subject", ""),
        "created_from": filtros["created_from"].isoformat() if filtros.get("created_from") else "",
        "created_to": filtros["created_to"].isoformat() if filtros.get("created_to") else "",
        "sort_by": filtros.get("sort_by", "updated_desc"),
    }

def validar_dados_geracao(
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str = "",
) -> dict:
    client_name = (client_name or "").strip()
    document_type = (document_type or "").strip()
    case_subject = (case_subject or "").strip()
    facts = (facts or "").strip()
    requests = (requests or "").strip()
    legal_basis = (legal_basis or "").strip()

    if not client_name:
        raise ValueError("Informe o nome do cliente.")
    if len(client_name) < 3:
        raise ValueError("O nome do cliente deve ter pelo menos 3 caracteres.")

    if not document_type:
        raise ValueError("Selecione o tipo de documento.")
    if document_type not in TIPOS_DE_DOCUMENTO:
        raise ValueError("Tipo de documento inválido.")

    if not case_subject:
        raise ValueError("Informe o assunto do caso.")
    if len(case_subject) < 5:
        raise ValueError("O assunto do caso deve ter pelo menos 5 caracteres.")

    if not facts:
        raise ValueError("Informe os fatos do caso.")
    if len(facts) < 20:
        raise ValueError("Os fatos do caso devem ter pelo menos 20 caracteres.")

    if not requests:
        raise ValueError("Informe os pedidos.")
    if len(requests) < 10:
        raise ValueError("Os pedidos devem ter pelo menos 10 caracteres.")

    return {
        "client_name": client_name,
        "document_type": document_type,
        "case_subject": case_subject,
        "facts": facts,
        "requests": requests,
        "legal_basis": legal_basis,
    }


def serializar_ids_documentos(document_ids: list[int] | None) -> str | None:
    ids = [str(document_id) for document_id in (document_ids or []) if isinstance(document_id, int)]
    return ",".join(ids) if ids else None


def desserializar_ids_documentos(source_document_ids: str | None) -> list[int]:
    if not source_document_ids:
        return []

    ids: list[int] = []
    for parte in source_document_ids.split(","):
        valor = parte.strip()
        if not valor:
            continue
        try:
            ids.append(int(valor))
        except ValueError:
            continue
    return ids


def criar_geracao(
    db: Session,
    user_id: int,
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str,
    context_used: str,
    generated_text: str,
    generation_strategy: str = "rule_based",
    llm_provider: str | None = None,
    llm_model: str | None = None,
    llm_response_id: str | None = None,
    llm_error: str | None = None,
    writing_profile_id: int | None = None,
    source_document_ids: str | None = None,
    tags: str | None = None,
    is_favorite: bool = False,
    status: str | None = None,
) -> Generation:
    if writing_profile_id is not None:
        perfil_valido = (
            db.query(WritingProfile)
            .filter(WritingProfile.id == writing_profile_id, WritingProfile.user_id == user_id)
            .first()
        )
        if not perfil_valido:
            writing_profile_id = None

    document_ids = desserializar_ids_documentos(source_document_ids)
    documentos_relacionados = []

    if document_ids:
        documentos_relacionados = (
            db.query(Document)
            .filter(Document.user_id == user_id, Document.id.in_(document_ids))
            .all()
        )
        documentos_por_id = {documento.id: documento for documento in documentos_relacionados}
        documentos_relacionados = [
            documentos_por_id[document_id]
            for document_id in document_ids
            if document_id in documentos_por_id
        ]

    nova_geracao = Generation(
        user_id=user_id,
        client_name=_normalizar_texto(client_name),
        document_type=_normalizar_texto(document_type),
        case_subject=_normalizar_texto(case_subject),
        facts=_normalizar_texto(facts),
        requests=_normalizar_texto(requests),
        legal_basis=_normalizar_texto(legal_basis),
        context_used=_normalizar_texto(context_used),
        generated_text=_normalizar_texto(generated_text),
        generation_strategy=_normalizar_texto(generation_strategy) or "rule_based",
        llm_provider=_normalizar_texto(llm_provider) or None,
        llm_model=_normalizar_texto(llm_model) or None,
        llm_response_id=_normalizar_texto(llm_response_id) or None,
        llm_error=_normalizar_texto(llm_error) or None,
        writing_profile_id=writing_profile_id,
        source_document_ids=serializar_ids_documentos([documento.id for documento in documentos_relacionados]),
        tags=_normalizar_texto(tags) or None,
        is_favorite=bool(is_favorite),
        status=_normalizar_texto(status) or None,
        updated_at=agora_brasil(),
    )
    nova_geracao.documents = documentos_relacionados

    db.add(nova_geracao)
    db.commit()
    db.refresh(nova_geracao)
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=nova_geracao.id,
        action="create",
        entity_version=nova_geracao.version,
        snapshot=serializar_entidade_para_auditoria(nova_geracao),
    )
    db.commit()
    return nova_geracao

def atualizar_geracao(
    db: Session,
    geracao: Generation,
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str,
    context_used: str,
    generated_text: str,
    generation_strategy: str = "rule_based",
    llm_provider: str | None = None,
    llm_model: str | None = None,
    llm_response_id: str | None = None,
    llm_error: str | None = None,
    writing_profile_id: int | None = None,
    source_document_ids: str | None = None,
    tags: str | None = None,
    status: str | None = None,
) -> Generation:
    if writing_profile_id is not None:
        perfil_valido = (
            db.query(WritingProfile)
            .filter(WritingProfile.id == writing_profile_id, WritingProfile.user_id == geracao.user_id)
            .first()
        )
        if not perfil_valido:
            writing_profile_id = None

    document_ids = desserializar_ids_documentos(source_document_ids)
    documentos_relacionados = []

    if document_ids:
        documentos_relacionados = (
            db.query(Document)
            .filter(Document.user_id == geracao.user_id, Document.id.in_(document_ids))
            .all()
        )
        documentos_por_id = {documento.id: documento for documento in documentos_relacionados}
        documentos_relacionados = [
            documentos_por_id[document_id]
            for document_id in document_ids
            if document_id in documentos_por_id
        ]

    geracao.client_name = _normalizar_texto(client_name)
    geracao.document_type = _normalizar_texto(document_type)
    geracao.case_subject = _normalizar_texto(case_subject)
    geracao.facts = _normalizar_texto(facts)
    geracao.requests = _normalizar_texto(requests)
    geracao.legal_basis = _normalizar_texto(legal_basis)
    geracao.context_used = _normalizar_texto(context_used)
    geracao.generated_text = _normalizar_texto(generated_text)
    geracao.generation_strategy = _normalizar_texto(generation_strategy) or "rule_based"
    geracao.llm_provider = _normalizar_texto(llm_provider) or None
    geracao.llm_model = _normalizar_texto(llm_model) or None
    geracao.llm_response_id = _normalizar_texto(llm_response_id) or None
    geracao.llm_error = _normalizar_texto(llm_error) or None
    geracao.writing_profile_id = writing_profile_id
    geracao.documents = documentos_relacionados
    geracao.source_document_ids = serializar_ids_documentos([documento.id for documento in documentos_relacionados])

    geracao.tags = _normalizar_texto(tags) or None
    geracao.status = _normalizar_texto(status) or None
    geracao.version = int(getattr(geracao, "version", 1) or 1) + 1
    geracao.updated_at = agora_brasil()

    db.add(geracao)
    db.commit()
    db.refresh(geracao)
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=geracao.id,
        action="update",
        entity_version=geracao.version,
        snapshot=serializar_entidade_para_auditoria(geracao),
    )
    db.commit()
    return geracao

def buscar_geracao_por_id(db: Session, generation_id: int, user_id: int) -> Generation | None:
    return (
        db.query(Generation)
        .options(joinedload(Generation.writing_profile), joinedload(Generation.documents))
        .filter(Generation.id == generation_id, Generation.user_id == user_id)
        .first()
    )


def listar_geracoes(
    db: Session,
    user_id: int,
    filtros: dict | None = None,
):
    filtros = filtros or {}

    query = (
        db.query(Generation)
        .options(joinedload(Generation.writing_profile), joinedload(Generation.documents))
        .filter(Generation.user_id == user_id)
    )

    search_term = filtros.get("search_term")
    if search_term:
        like_term = f"%{search_term}%"
        query = query.filter(
            or_(
                Generation.client_name.ilike(like_term),
                Generation.document_type.ilike(like_term),
                Generation.case_subject.ilike(like_term),
                Generation.facts.ilike(like_term),
                Generation.requests.ilike(like_term),
                Generation.legal_basis.ilike(like_term),
                Generation.generated_text.ilike(like_term),
                Generation.tags.ilike(like_term),
                Generation.status.ilike(like_term),
            )
        )

    document_type = filtros.get("document_type")
    if document_type:
        query = query.filter(Generation.document_type == document_type)

    writing_profile_id = filtros.get("writing_profile_id")
    if writing_profile_id:
        query = query.filter(Generation.writing_profile_id == writing_profile_id)

    if filtros.get("sem_perfil"):
        query = query.filter(Generation.writing_profile_id.is_(None))

    client_name = filtros.get("client_name")
    if client_name:
        query = query.filter(Generation.client_name.ilike(f"%{client_name}%"))

    case_subject = filtros.get("case_subject")
    if case_subject:
        query = query.filter(Generation.case_subject.ilike(f"%{case_subject}%"))

    created_from: date | None = filtros.get("created_from")
    if created_from:
        query = query.filter(Generation.created_at >= datetime.combine(created_from, time.min))

    created_to: date | None = filtros.get("created_to")
    if created_to:
        query = query.filter(Generation.created_at <= datetime.combine(created_to, time.max))

    sort_by = filtros.get("sort_by", "updated_desc")

    if sort_by == "updated_asc":
        query = query.order_by(
            Generation.is_pinned.desc(),
            Generation.is_favorite.desc(),
            Generation.updated_at.asc(),
            Generation.id.asc(),
        )
    elif sort_by == "created_desc":
        query = query.order_by(
            Generation.is_pinned.desc(),
            Generation.is_favorite.desc(),
            Generation.created_at.desc(),
            Generation.id.desc(),
        )
    elif sort_by == "created_asc":
        query = query.order_by(
            Generation.is_pinned.desc(),
            Generation.is_favorite.desc(),
            Generation.created_at.asc(),
            Generation.id.asc(),
        )
    elif sort_by == "client_asc":
        query = query.order_by(
            Generation.is_pinned.desc(),
            Generation.is_favorite.desc(),
            Generation.client_name.asc(),
            Generation.id.asc(),
        )
    elif sort_by == "client_desc":
        query = query.order_by(
            Generation.is_pinned.desc(),
            Generation.is_favorite.desc(),
            Generation.client_name.desc(),
            Generation.id.desc(),
        )
    else:
        query = query.order_by(
            Generation.is_pinned.desc(),
            Generation.is_favorite.desc(),
            Generation.updated_at.desc(),
            Generation.id.desc(),
        )

    return query.all()


def montar_resumo_geracao(geracao: Generation) -> dict:
    perfil = geracao.writing_profile
    document_ids = geracao.document_ids

    return {
        "id": geracao.id,
        "client_name": geracao.client_name,
        "document_type": geracao.document_type,
        "case_subject": geracao.case_subject,
        "facts": geracao.facts,
        "requests": geracao.requests,
        "legal_basis": geracao.legal_basis,
        "context_used": geracao.context_used,
        "generated_text": geracao.generated_text,
        "generation_strategy": getattr(geracao, "generation_strategy", "rule_based") or "rule_based",
        "llm_provider": getattr(geracao, "llm_provider", None),
        "llm_model": getattr(geracao, "llm_model", None),
        "llm_response_id": getattr(geracao, "llm_response_id", None),
        "llm_error": getattr(geracao, "llm_error", None),
        "source_document_ids": geracao.source_document_ids,
        "writing_profile_id": geracao.writing_profile_id,
        "writing_profile_name": perfil.profile_name if perfil else "Sem perfil",
        "tags": geracao.tags or "",
        "status": geracao.status or "",
        "is_pinned": bool(geracao.is_pinned),
        "is_favorite": bool(geracao.is_favorite),
        "created_at": geracao.created_at,
        "updated_at": geracao.updated_at,
        "document_count": len(document_ids),
        "facts_preview": resumir_texto(geracao.facts, 180),
        "requests_preview": resumir_texto(geracao.requests, 180),
        "generated_text_preview": resumir_texto(geracao.generated_text, 220),
    }

def alternar_fixacao_geracao(db: Session, geracao: Generation) -> Generation:
    geracao.is_pinned = not bool(geracao.is_pinned)
    geracao.version = int(getattr(geracao, "version", 1) or 1) + 1
    geracao.updated_at = agora_brasil()

    db.add(geracao)
    db.commit()
    db.refresh(geracao)
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=geracao.id,
        action="toggle_pin",
        entity_version=geracao.version,
        snapshot=serializar_entidade_para_auditoria(geracao),
    )
    db.commit()
    return geracao


def alternar_favorito_geracao(db: Session, geracao: Generation) -> Generation:
    geracao.is_favorite = not bool(geracao.is_favorite)
    geracao.version = int(getattr(geracao, "version", 1) or 1) + 1
    geracao.updated_at = agora_brasil()

    db.add(geracao)
    db.commit()
    db.refresh(geracao)
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=geracao.id,
        action="toggle_favorite",
        entity_version=geracao.version,
        snapshot=serializar_entidade_para_auditoria(geracao),
    )
    db.commit()
    return geracao

def excluir_geracao(db: Session, geracao: Generation) -> None:
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=geracao.id,
        action="delete",
        entity_version=int(getattr(geracao, "version", 1) or 1),
        snapshot=serializar_entidade_para_auditoria(geracao),
    )
    db.delete(geracao)
    db.commit()

def duplicar_geracao(
    db: Session,
    geracao_origem: Generation,
) -> Generation:
    nova_geracao = Generation(
        user_id=geracao_origem.user_id,
        client_name=geracao_origem.client_name,
        document_type=geracao_origem.document_type,
        case_subject=geracao_origem.case_subject,
        facts=geracao_origem.facts,
        requests=geracao_origem.requests,
        legal_basis=geracao_origem.legal_basis,
        context_used=geracao_origem.context_used,
        generated_text=geracao_origem.generated_text,
        writing_profile_id=geracao_origem.writing_profile_id,
        source_document_ids=serializar_ids_documentos(geracao_origem.document_ids),
        is_pinned=False,
        updated_at=agora_brasil(),
    )
    nova_geracao.documents = list(geracao_origem.documents)

    db.add(nova_geracao)
    db.commit()
    db.refresh(nova_geracao)
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=nova_geracao.id,
        action="duplicate",
        entity_version=nova_geracao.version,
        snapshot=serializar_entidade_para_auditoria(nova_geracao),
    )
    db.commit()
    return nova_geracao


def salvar_texto_geracao(db: Session, geracao: Generation, generated_text: str) -> Generation:
    geracao.generated_text = _normalizar_texto(generated_text)
    geracao.version = int(getattr(geracao, "version", 1) or 1) + 1
    geracao.updated_at = agora_brasil()

    db.add(geracao)
    db.commit()
    db.refresh(geracao)
    registrar_evento_auditoria(
        db,
        entity_type="generation",
        entity_id=geracao.id,
        action="save_text",
        entity_version=geracao.version,
        snapshot=serializar_entidade_para_auditoria(geracao),
    )
    db.commit()
    return geracao


def aplicar_template_juridico_pronto(document_type: str) -> dict:
    tipo = (document_type or "").strip()

    if not tipo:
        raise ValueError("Informe o tipo de documento para aplicar um template.")

    template = TEMPLATES_JURIDICOS_PRONTOS.get(tipo)
    if not template:
        raise ValueError("Não existe template pronto para este tipo de documento.")

    return {
        "document_type": tipo,
        "case_subject": template["case_subject"],
        "facts": template["facts"],
        "requests": template["requests"],
        "legal_basis": template["legal_basis"],
        "template_title": template["titulo"],
        "template_description": template["descricao"],
    }


def montar_contexto_documentos(documentos_selecionados: list) -> str:
    if not documentos_selecionados:
        return "Nenhum documento base selecionado."

    blocos = []
    for documento in documentos_selecionados:
        resumo = (getattr(documento, "summary", "") or "").strip()
        texto_extraido = (getattr(documento, "extracted_text", "") or "").strip()
        nome_arquivo = getattr(documento, "original_filename", "Documento sem nome")
        tipo_arquivo = getattr(documento, "file_type", "arquivo")
        documento_id = getattr(documento, "id", "?")

        conteudo = resumo or texto_extraido or "Documento sem conteúdo textual extraído."
        if len(conteudo) > 600:
            conteudo = conteudo[:600].rstrip() + "..."

        blocos.append(
            f"Documento #{documento_id} - {nome_arquivo} ({tipo_arquivo})\n"
            f"{conteudo}"
        )

    return "\n" + ("\n" + ("-" * 60) + "\n").join(blocos)


def montar_contexto_inteligente(
    *,
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str,
    writing_profile=None,
    documentos_selecionados: list | None = None,
) -> str:
    return build_smart_context(
        client_name=client_name,
        document_type=document_type,
        case_subject=case_subject,
        facts=facts,
        requests=requests,
        legal_basis=legal_basis,
        writing_profile=writing_profile,
        documentos_selecionados=documentos_selecionados or [],
    )


def montar_contexto_perfil_escrita(profile) -> str:
    if not profile:
        return "Nenhum perfil de escrita selecionado."

    expressoes = (profile.recurring_expressions or "").strip()
    observacoes = (profile.legal_style_notes or "").strip()

    return (
        f"Perfil de escrita selecionado:\n"
        f"- Nome do perfil: {profile.profile_name}\n"
        f"- Advogado: {profile.lawyer_name or 'Não informado'}\n"
        f"- Escritório: {profile.office_name or 'Não informado'}\n"
        f"- Tom: {profile.tone or 'Formal'}\n"
        f"- Estilo de qualificação: {profile.qualification_style or 'Não informado'}\n"
        f"- Frase de abertura: {profile.opening_phrase or 'Não informada'}\n"
        f"- Introdução dos pedidos: {profile.request_intro or 'Não informada'}\n"
        f"- Frase de fechamento: {profile.closing_phrase or 'Não informada'}\n"
        f"- Observações de estilo: {observacoes or 'Não informado'}\n"
        f"- Expressões recorrentes: {expressoes or 'Não informado'}"
    )


def _tipo_normalizado(document_type: str) -> str:
    return (document_type or "").strip().lower()


def _limpar_pontuacao_final(texto: str) -> str:
    return (texto or "").strip().rstrip(".;:,")


def _primeira_frase(texto: str, limite: int = 220) -> str:
    texto_limpo = " ".join((texto or "").split())

    if not texto_limpo:
        return ""

    for separador in [". ", "; ", ": "]:
        if separador in texto_limpo:
            trecho = texto_limpo.split(separador)[0].strip()
            if trecho:
                return trecho[:limite].rstrip(".,;:") + "."

    if len(texto_limpo) <= limite:
        return texto_limpo.rstrip(".,;:") + "."

    return texto_limpo[:limite].rstrip(".,;:") + "..."


def _texto_em_linhas(texto: str) -> list[str]:
    linhas = []
    for linha in (texto or "").splitlines():
        linha_limpa = linha.strip()
        if linha_limpa:
            linhas.append(linha_limpa)
    return linhas


def _bulletizar_texto(texto: str) -> str:
    linhas = _texto_em_linhas(texto)

    if not linhas:
        texto_limpo = (texto or "").strip()
        return texto_limpo

    bullets = []
    for linha in linhas:
        linha_sem_prefixo = linha.lstrip("-•* ").strip()
        if linha_sem_prefixo:
            bullets.append(f"- {linha_sem_prefixo}")

    return "\n".join(bullets)


def _coletar_trechos_documentais(documentos_selecionados: list | None, limite: int = 3) -> list[str]:
    if not documentos_selecionados:
        return []

    trechos = []
    for documento in documentos_selecionados[:limite]:
        resumo = (getattr(documento, "summary", "") or "").strip()
        texto_extraido = (getattr(documento, "extracted_text", "") or "").strip()
        nome_arquivo = getattr(documento, "original_filename", "Documento sem nome")

        base = resumo or texto_extraido
        if not base:
            continue

        frase = _primeira_frase(base, 220)
        if frase:
            trechos.append(f"- {nome_arquivo}: {frase}")

    return trechos


def _montar_qualificacao(profile, client_name: str) -> str:
    if profile and (profile.qualification_style or "").strip():
        estilo = profile.qualification_style.strip()
        return estilo.replace("{cliente}", client_name)

    return f"{client_name}, já devidamente qualificado(a), vem, respeitosamente, à presença de Vossa Excelência, propor a presente, nos termos a seguir expostos:"


def _montar_fundamentacao(
    document_type: str,
    case_subject: str,
    legal_basis: str,
    documentos_selecionados: list | None = None,
) -> str:
    base_juridica = (legal_basis or "").strip()
    assunto = _limpar_pontuacao_final(case_subject)

    trechos_documentais = _coletar_trechos_documentais(documentos_selecionados)

    if _tipo_normalizado(document_type) == "contrato":
        partes = [
            "A presente minuta contratual observa os princípios da autonomia privada, da boa-fé objetiva, da função social do contrato e da força obrigatória das convenções.",
        ]

        if base_juridica:
            partes.append(
                f"Considera-se, ainda, como base jurídica inicial, em tese, o seguinte: {base_juridica}"
            )
        else:
            partes.append(
                "Devem ser observadas as normas do Código Civil aplicáveis ao objeto contratado, bem como cláusulas essenciais de equilíbrio, segurança jurídica e executabilidade."
            )

        if assunto:
            partes.append(
                f"As cláusulas contratuais devem refletir com precisão o objeto relacionado a {assunto}, delimitando obrigações, responsabilidades, prazos, forma de pagamento e hipóteses de rescisão."
            )

        if trechos_documentais:
            partes.append("Os documentos base reforçam os seguintes pontos relevantes:")
            partes.extend(trechos_documentais)

        return "\n\n".join(partes)

    partes = []

    if base_juridica:
        partes.append(
            "A pretensão deduzida encontra amparo, em tese, nos seguintes fundamentos jurídicos:\n\n"
            f"{base_juridica}"
        )
    else:
        partes.append(
            "A pretensão deduzida encontra amparo, em tese, na legislação aplicável ao caso concreto, "
            "nos princípios da boa-fé, da razoabilidade, da efetividade da tutela jurisdicional e nas "
            "demais normas pertinentes à controvérsia apresentada."
        )

    if assunto:
        partes.append(
            f"Em especial, o enquadramento jurídico deve considerar a controvérsia relacionada a {assunto}, "
            "com análise da responsabilidade aplicável, da adequação da medida pretendida e da coerência entre fatos, fundamentos e pedidos."
        )

    if trechos_documentais:
        partes.append("Os documentos base selecionados reforçam, em síntese, os seguintes elementos contextuais:")
        partes.extend(trechos_documentais)

    return "\n\n".join(partes)


def _montar_pedidos(document_type: str, requests: str, profile=None) -> str:
    pedidos_formatados = _bulletizar_texto(requests)
    introducao = ""

    if profile and (profile.request_intro or "").strip():
        introducao = profile.request_intro.strip()
    else:
        tipo = _tipo_normalizado(document_type)
        if tipo == "contestação":
            introducao = "Diante do exposto, requer:"
        elif tipo == "réplica":
            introducao = "Diante do exposto, requer:"
        elif tipo == "manifestação":
            introducao = "Diante do exposto, requer:"
        elif tipo == "parecer jurídico":
            introducao = "Com base na análise realizada, conclui-se e recomenda-se:"
        elif tipo == "contrato":
            introducao = "Ficam estabelecidas as seguintes disposições essenciais:"
        elif tipo == "notificação extrajudicial":
            introducao = "Diante do exposto, fica a parte notificada cientificada e intimada para:"
        elif tipo == "recurso":
            introducao = "Diante do exposto, requer:"
        else:
            introducao = "Diante do exposto, requer:"

    if not pedidos_formatados:
        return introducao

    return f"{introducao}\n\n{pedidos_formatados}"


def _montar_fechamento(document_type: str, profile=None) -> str:
    if profile and (profile.closing_phrase or "").strip():
        return profile.closing_phrase.strip()

    tipo = _tipo_normalizado(document_type)
    if tipo == "parecer jurídico":
        return "É o parecer, s.m.j."
    if tipo == "contrato":
        return "Por estarem justas e contratadas, as partes firmam a presente minuta para os devidos fins."
    if tipo == "notificação extrajudicial":
        return "Sem mais para o momento, aguarda-se o cumprimento da providência ora exigida."
    return "Termos em que,\nPede deferimento."


def _montar_assinatura(profile) -> str:
    if not profile:
        return ""

    partes = []
    if (profile.lawyer_name or "").strip():
        partes.append(profile.lawyer_name.strip())
    if (profile.office_name or "").strip():
        partes.append(profile.office_name.strip())

    return "\n".join(partes).strip()


def _gerar_rascunho_local(
    *,
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str,
    context_used: str,
    writing_profile=None,
    documentos_selecionados: list | None = None,
) -> str:
    tipo_normalizado = _tipo_normalizado(document_type)

    titulo = document_type.upper().strip()
    qualificacao = _montar_qualificacao(writing_profile, client_name)
    fundamentacao = _montar_fundamentacao(
        document_type=document_type,
        case_subject=case_subject,
        legal_basis=legal_basis,
        documentos_selecionados=documentos_selecionados,
    )
    pedidos = _montar_pedidos(document_type, requests, writing_profile)
    fechamento = _montar_fechamento(document_type, writing_profile)
    assinatura = _montar_assinatura(writing_profile)

    assunto_limpo = _limpar_pontuacao_final(case_subject)
    fatos_iniciais = (facts or "").strip()

    if tipo_normalizado == "contrato":
        texto = f"""
{titulo}

As partes interessadas resolvem celebrar a presente minuta contratual, observadas as disposições a seguir:

I - DO OBJETO

A presente minuta tem por objeto disciplinar, em linhas gerais, a relação jurídica relacionada a {assunto_limpo or 'objeto a ser especificado pelas partes'}, conforme contexto apresentado.

II - DO CONTEXTO CONTRATUAL

{fatos_iniciais}

III - DAS BASES JURÍDICAS E DIRETRIZES

{fundamentacao}

IV - DAS CLÁUSULAS ESSENCIAIS

{pedidos}

V - DAS DISPOSIÇÕES FINAIS

{fechamento}
""".strip()

        if assinatura:
            texto += f"\n\n{assinatura}"

        return texto

    titulo_fatos = "I - DOS FATOS"
    titulo_fundamentacao = "II - DA FUNDAMENTAÇÃO JURÍDICA"
    titulo_final = "III - DOS PEDIDOS"

    introducao_final = pedidos
    conteudo_final = ""
    fechamento_final = fechamento

    if tipo_normalizado == "contestação":
        titulo_fatos = "I - SÍNTESE DA DEMANDA"
        titulo_fundamentacao = "II - DOS FUNDAMENTOS DE DEFESA"
        titulo_final = "III - DOS REQUERIMENTOS FINAIS"
    elif tipo_normalizado == "réplica":
        titulo_fatos = "I - DA SÍNTESE DA CONTESTAÇÃO"
        titulo_fundamentacao = "II - DA IMPUGNAÇÃO AOS ARGUMENTOS DEFENSIVOS"
        titulo_final = "III - DOS REQUERIMENTOS"
    elif tipo_normalizado == "manifestação":
        titulo_fatos = "I - DO CONTEXTO PROCESSUAL"
        titulo_fundamentacao = "II - DAS CONSIDERAÇÕES JURÍDICAS"
        titulo_final = "III - DO REQUERIMENTO"
    elif tipo_normalizado == "parecer jurídico":
        titulo_fatos = "I - DO RELATÓRIO"
        titulo_fundamentacao = "II - DA ANÁLISE JURÍDICA"
        titulo_final = "III - DA CONCLUSÃO"
    elif tipo_normalizado == "notificação extrajudicial":
        titulo_fatos = "I - DOS FATOS"
        titulo_fundamentacao = "II - DOS FUNDAMENTOS"
        titulo_final = "III - DA NOTIFICAÇÃO E PROVIDÊNCIAS EXIGIDAS"
    elif tipo_normalizado == "recurso":
        titulo_fatos = "I - DA SÍNTESE DA DECISÃO RECORRIDA"
        titulo_fundamentacao = "II - DAS RAZÕES RECURSAIS"
        titulo_final = "III - DOS PEDIDOS RECURSAIS"

    texto = f"""
{titulo}

{qualificacao}

A presente demanda decorre de {assunto_limpo or 'questão jurídica submetida à análise'}, conforme fatos e fundamentos a seguir expostos.

{titulo_fatos}

{fatos_iniciais}

{titulo_fundamentacao}

{fundamentacao}

{titulo_final}

{introducao_final}

{conteudo_final}

{fechamento_final}
""".strip()

    if assinatura:
        texto += f"\n\n{assinatura}"

    return texto


def gerar_rascunho_juridico(
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str,
    context_used: str,
    writing_profile=None,
    documentos_selecionados: list | None = None,
) -> str:
    return gerar_rascunho_juridico_com_metadata(
        client_name=client_name,
        document_type=document_type,
        case_subject=case_subject,
        facts=facts,
        requests=requests,
        legal_basis=legal_basis,
        context_used=context_used,
        writing_profile=writing_profile,
        documentos_selecionados=documentos_selecionados,
    ).text


def gerar_rascunho_juridico_com_metadata(
    client_name: str,
    document_type: str,
    case_subject: str,
    facts: str,
    requests: str,
    legal_basis: str,
    context_used: str,
    writing_profile=None,
    documentos_selecionados: list | None = None,
) -> AIGenerationResult:
    prompt_payload = build_advanced_prompt(
        client_name=client_name,
        document_type=document_type,
        case_subject=case_subject,
        facts=facts,
        requests=requests,
        legal_basis=legal_basis,
        smart_context=context_used,
        writing_profile=writing_profile,
        documentos_selecionados=documentos_selecionados or [],
    )

    return gerar_resultado_juridico_com_fallback(
        prompt_payload=prompt_payload,
        fallback_generator=lambda: _gerar_rascunho_local(
            client_name=client_name,
            document_type=document_type,
            case_subject=case_subject,
            facts=facts,
            requests=requests,
            legal_basis=legal_basis,
            context_used=context_used,
            writing_profile=writing_profile,
            documentos_selecionados=documentos_selecionados or [],
        ),
    )


def gerar_docx_da_geracao(geracao: Generation) -> bytes:
    documento = DocxDocument()

    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Times New Roman"
    estilo_normal.font.size = Pt(12)

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run((geracao.document_type or "MINUTA JURÍDICA").upper())
    run_titulo.bold = True

    documento.add_paragraph("")

    for bloco in (geracao.generated_text or "").split("\n\n"):
        bloco_limpo = bloco.strip()
        if not bloco_limpo:
            continue

        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafo.add_run(bloco_limpo)

    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def gerar_txt_da_geracao(geracao: Generation) -> bytes:
    return (geracao.generated_text or "").encode("utf-8")


def obter_datas_dashboard(db: Session) -> dict:
    hoje = agora_brasil().date()
    inicio_hoje = datetime.combine(hoje, time.min)
    fim_hoje = datetime.combine(hoje, time.max)

    inicio_7_dias = datetime.combine(hoje - timedelta(days=6), time.min)
    inicio_30_dias = datetime.combine(hoje - timedelta(days=29), time.min)
    inicio_mes = datetime.combine(hoje.replace(day=1), time.min)

    return {
        "hoje": hoje,
        "inicio_hoje": inicio_hoje,
        "fim_hoje": fim_hoje,
        "inicio_7_dias": inicio_7_dias,
        "inicio_30_dias": inicio_30_dias,
        "inicio_mes": inicio_mes,
    }


def obter_resumo_dashboard_geracoes(db: Session, user_id: int) -> dict:
    datas = obter_datas_dashboard(db)

    total = db.query(Generation).filter(Generation.user_id == user_id).count()

    total_hoje = (
        db.query(Generation)
        .filter(Generation.user_id == user_id)
        .filter(Generation.created_at >= datas["inicio_hoje"])
        .filter(Generation.created_at <= datas["fim_hoje"])
        .count()
    )

    total_7_dias = (
        db.query(Generation)
        .filter(Generation.user_id == user_id)
        .filter(Generation.created_at >= datas["inicio_7_dias"])
        .count()
    )

    total_30_dias = (
        db.query(Generation)
        .filter(Generation.user_id == user_id)
        .filter(Generation.created_at >= datas["inicio_30_dias"])
        .count()
    )

    total_mes = (
        db.query(Generation)
        .filter(Generation.user_id == user_id)
        .filter(Generation.created_at >= datas["inicio_mes"])
        .count()
    )

    total_fixadas = (
        db.query(Generation)
        .filter(Generation.user_id == user_id)
        .filter(Generation.is_pinned.is_(True))
        .count()
    )

    return {
        "total": total,
        "total_hoje": total_hoje,
        "total_7_dias": total_7_dias,
        "total_30_dias": total_30_dias,
        "total_mes": total_mes,
        "total_fixadas": total_fixadas,
    }


def obter_ultimas_geracoes(db: Session, user_id: int, limite: int = 5) -> list[Generation]:
    return (
        db.query(Generation)
        .options(joinedload(Generation.writing_profile))
        .filter(Generation.user_id == user_id)
        .order_by(Generation.updated_at.desc(), Generation.id.desc())
        .limit(limite)
        .all()
    )
